"""Chat orchestration: manager decomposes, skill agents collaborate, reply as one agent."""

from __future__ import annotations

import json
import re

from backend.config import AGENT_OUTPUT_FOLDER
from backend.services.agent_builder import agent_file_slug
from backend.services.agent_instructions import build_agent_instruction_block, prepend_agent_instructions
from backend.services.chat_progress import ChatGenerationCancelled, ChatProgressReporter
from backend.services.conversation_memory import ConversationMemoryContext
from backend.services.cursor_llm import complete as cursor_complete
from backend.services.cursor_llm import chat as cursor_chat
from backend.services.cursor_llm import is_configured as cursor_is_configured
from backend.services.local_replies import normalize_user_query, try_instant_reply
from backend.services.skill_framework import MANAGER_ID, PIPELINE_ORDER


def _compose_system(
    system: str,
    agent_context: dict | None,
    memory_context: ConversationMemoryContext | None,
) -> str:
    """Agent settings + retrieved conversation memory on every LLM call."""
    composed = prepend_agent_instructions(system, agent_context)
    if memory_context and memory_context.memory_prompt_block:
        composed = f"{composed}\n\n{memory_context.memory_prompt_block}"
    return composed


def _wrap_user_message_for_follow_up(
    user_message: str,
    cloud_agent_id: str | None,
    agent_context: dict | None,
    memory_context: ConversationMemoryContext | None,
    recent_history: list[dict] | None = None,
) -> str:
    """Re-inject settings + memory on Cursor follow-ups (system prompt is not resent)."""
    if not cloud_agent_id:
        return user_message
    parts: list[str] = []
    instr = build_agent_instruction_block(agent_context)
    if instr:
        parts.append(instr)
    memory_block = (memory_context.memory_prompt_block if memory_context else "") or ""
    if memory_block:
        parts.append(memory_block)
    elif recent_history:
        from backend.services.conversation_memory import format_recent_conversation_block

        recent_block = format_recent_conversation_block(recent_history)
        if recent_block:
            parts.append(recent_block)
    parts.append(f"CURRENT REQUEST:\n{user_message}")
    return "\n\n".join(parts)


def _effective_history(
    history: list[dict],
    memory_context: ConversationMemoryContext | None,
) -> list[dict]:
    if memory_context and memory_context.short_term_history:
        return [
            {"role": h.get("role"), "content": h.get("content")}
            for h in memory_context.short_term_history
            if h.get("role") in ("user", "assistant")
        ]
    return [h for h in history if h.get("role") in ("user", "assistant")]

TRIVIAL_MESSAGES = frozenset({
    "hi", "hello", "hey", "thanks", "thank you", "ok", "okay", "yes", "no", "help",
})

COMPLEX_KEYWORDS = (
    "analyze", "analysis", "report", "deck", "slide", "funnel", "forecast",
    "investigate", "compare", "summarize", "summary", "breakdown", "audit",
    "recommend", "migration", "dashboard", "variance", "prepare", "build",
)

# Keywords that suggest a specialist skill is needed (not exhaustive — manager triage refines).
SKILL_KEYWORDS: dict[str, tuple[str, ...]] = {
    "sql query": ("sql", "query", "database", "warehouse", "table", "metric", "kpi", "funnel", "cohort", "pull data"),
    "python automation": ("python", "script", "automate", "pipeline", "etl"),
    "statistical modeling": ("regression", "statistic", "variance", "predict", "statistical model"),
    "financial analysis": ("financial", "revenue", "p&l", "margin", "fp&a", "kpi trend"),
    "data visualization": ("chart", "graph", "visualiz", "dashboard", "tableau", "plot"),
    "process optimization": ("process", "workflow", "efficiency", "bottleneck", "optimize"),
    "business acumen": ("strategy", "market", "stakeholder", "business case", "go to market", "gtm", "implication", "impact", "investment thesis"),
    "growth marketing": ("marketing", "campaign", "acquisition", "conversion", "growth"),
    "compliance & governance": ("compliance", "regulatory", "governance", "audit trail", "policy", "regulation", "legal", "eu ai act", "ai act"),
    "presentation & storytelling": ("presentation", "deck", "slides", "storytell", "pitch"),
    "written communication": ("memo", "email", "write up", "write-up", "narrative", "executive summary"),
    "domain expertise": ("industry", "domain", "sector", "vertical", "macro", "landscape"),
    "scenario planning": ("scenario", "stress test", "what-if", "what if", "downside case", "upside case", "sensitivity"),
    "valuation": ("valuation", "dcf", "discounted cash", "comps", "multiple", "enterprise value", "fair value"),
    "financial modeling": ("financial model", "spreadsheet model", "three-statement", "build a model", "model assumptions"),
    "forecasting": ("forecast", "projection", "outlook", "project revenue", "project earnings"),
    "cash flow": ("cash flow", "fcf", "free cash", "liquidity", "runway", "working capital"),
    "investor reporting": ("investor", "lp letter", "portfolio", "allocation", "investment memo", "ic memo"),
    "catalyst hunter": ("catalyst", "event risk", "earnings date", "timing", "near-term driver"),
    "financial accounting": (
        "p&l", "profit and loss", "cash flow", "cashflow", "fp&a", "variance", "budget",
        "accounting", "balance sheet", "income statement", "excel", "general ledger",
    ),
    "financial modeling & valuation": (
        "valuation", "dcf", "financial model", "forecast", "scenario", "comps", "sensitivity",
        "enterprise value", "three-statement",
    ),
    "investor & board reporting": (
        "investor", "board", "lp letter", "ic memo", "investment memo", "executive summary",
    ),
    "data & analytics": (
        "sql", "query", "database", "python", "tableau", "dashboard", "warehouse", "etl", "analytics",
    ),
    "statistical analysis": ("regression", "statistic", "predict", "machine learning", "correlation"),
}

# Task intents detected from the request (semantic layer — not keyword→skill mapping).
MACRO_OUTLOOK_TERMS = (
    "recession", "inflation", "unemployment", "yield curve", "gdp", "cpi", "ppi",
    "interest rate", "monetary", "federal reserve", "fed funds", "stagflation",
    "business cycle", "economic outlook", "soft landing", "hard landing",
)
STRATEGY_RECOMMENDATION_TERMS = (
    "defensive", "positioning", "allocation", "portfolio", "recommend", "suggest",
    "overweight", "underweight", "hedge", "tactical", "sector rotation", "risk-off",
    "risk-on", "asset mix", "rebalance",
)
DELIVERABLE_REQUEST_TERMS = (
    "report", "memo", "deck", "slides", "board", "executive summary", "write up",
    "write-up", "presentation", "deliverable",
)
EXECUTIVE_WORKFLOW_TERMS = (
    "methodology", "recommendation", "rigorous", "executive", "board-ready",
)
DATA_PULL_TERMS = (
    "sql", "query", "database", "warehouse", "pull data", "extract", "table",
)

# Intent → skill substring weights (semantic routing; complements keyword hints).
INTENT_SKILL_WEIGHTS: dict[str, list[tuple[str, int]]] = {
    "macro_outlook": [
        ("business acumen", 8),
        ("domain expertise", 7),
        ("scenario planning", 8),
        ("investor", 4),
        ("statistical modeling", -4),
        ("valuation", -6),
        ("financial modeling", -5),
    ],
    "strategy_recommendation": [
        ("business acumen", 9),
        ("scenario planning", 6),
        ("investor", 4),
        ("domain expertise", 5),
        ("statistical modeling", -3),
        ("board reporting", -4),
    ],
    "qualitative_forecast": [
        ("scenario planning", 7),
        ("business acumen", 6),
        ("domain expertise", 5),
        ("forecasting", 3),
        ("statistical modeling", 0),
    ],
    "regulatory_research": [
        ("compliance", 8),
        ("governance", 8),
        ("business acumen", 5),
        ("domain expertise", 4),
        ("investor", 3),
        ("valuation", -6),
        ("financial modeling", -6),
    ],
    "quant_modeling": [
        ("financial modeling", 8),
        ("valuation", 7),
        ("statistical modeling", 6),
        ("forecasting", 5),
        ("cash flow", 5),
        ("scenario planning", 4),
    ],
    "data_pull": [
        ("sql", 8),
        ("data & analytics", 7),
        ("python automation", 6),
    ],
    "stakeholder_delivery": [
        ("investor", 5),
        ("board reporting", 5),
        ("written communication", 5),
        ("presentation", 5),
    ],
    "executive_workflow": [
        ("presentation", 7),
        ("written communication", 7),
        ("financial analysis", 6),
        ("statistical modeling", 5),
        ("business acumen", 6),
        ("investor", 5),
    ],
}

LLM_RELEVANCE_THRESHOLD = 4

# Skills that imply spreadsheet/valuation work — only delegate when the user asks for quant output.
MODELING_SKILL_MARKERS = (
    "valuation", "dcf", "financial modeling", "cash flow", "forecasting", "scenario planning",
)

RESEARCH_SYNTHESIS_SIGNALS = (
    "impact", "implication", "overview", "landscape", "how does", "how will", "what is the effect",
    "regulation", "regulatory", "policy", "macro", "geopolitical", "explain", "assess", "evaluate",
    "in general", "broadly", "outlook", "ramifications", "consequences",
)

QUANT_MODELING_SIGNALS = (
    "dcf", "discounted cash", "financial model", "build a model", "cash flow model",
    "valuation model", "sensitivity analysis", "scenario model", "forecast model",
    "project revenue", "three-statement", "spreadsheet", "quantify", "run the numbers",
    "build a forecast", "model the", "value the company", "enterprise value",
)

DIRECT_ANSWER_PATTERNS = (
    r"\bwhat(?:'s| is) the date\b",
    r"\bwhat date\b",
    r"\bwhat(?:'s| is) the time\b",
    r"\bwhat time\b",
    r"\bwhat day is\b",
    r"\btoday(?:'s)? date\b",
    r"\bcurrent date\b",
    r"\bwhat does .+ mean\b",
    r"\bwhat is .+\?\s*$",
    r"\bwho are you\b",
    r"\bwhat can you do\b",
    r"\bhow are you\b",
)

# Maps skill categories to how the manager assigns work (not the raw user prompt).
SKILL_ROLE: dict[str, str] = {
    "sql query": "data_gathering",
    "python automation": "data_gathering",
    "statistical modeling": "analysis",
    "financial analysis": "analysis",
    "data visualization": "visualization",
    "process optimization": "analysis",
    "business acumen": "interpretation",
    "growth marketing": "strategy",
    "compliance & governance": "review",
    "presentation & storytelling": "delivery",
    "written communication": "delivery",
    "domain expertise": "interpretation",
    "financial accounting": "analysis",
    "financial modeling & valuation": "analysis",
    "investor & board reporting": "delivery",
    "data & analytics": "data_gathering",
    "statistical analysis": "analysis",
}

DEFAULT_MAX_SUBTASKS = 3
CONFIRMED_MAX_SUBTASKS_CEILING = 6


def _max_subtasks(skill_agents: list[dict], *, expanded: bool = False) -> int:
    """Lean cap (3) for auto-run; up to 6 specialists after expanded planning / approval."""
    if not expanded:
        return DEFAULT_MAX_SUBTASKS
    return min(len(skill_agents), CONFIRMED_MAX_SUBTASKS_CEILING)


def _use_expanded_planning_cap(
    user_message: str,
    skill_agents: list[dict],
    matched_skills: list[dict],
    user_feedback: str | None,
) -> bool:
    """Use the higher cap only when the user explicitly needs many distinct deliverables."""
    if user_feedback:
        return True
    if len(matched_skills) > DEFAULT_MAX_SUBTASKS:
        return True
    if _explicit_multi_deliverable_request(user_message):
        return True
    return False


def _requires_quantitative_modeling(user_message: str) -> bool:
    lower = user_message.lower()
    return any(signal in lower for signal in QUANT_MODELING_SIGNALS)


def _is_research_synthesis_request(user_message: str) -> bool:
    lower = user_message.lower()
    return any(signal in lower for signal in RESEARCH_SYNTHESIS_SIGNALS)


def _explicit_multi_deliverable_request(user_message: str) -> bool:
    """True when the user names multiple distinct outputs (not just a long question)."""
    lower = user_message.lower()
    if not any(p in lower for p in (" and ", " then ", " also ", " plus ", " including ")):
        return False
    action_verbs = (
        "analyze", "build", "prepare", "create", "forecast", "model", "report",
        "recommend", "deliverable", "dashboard", "deck", "query", "pull", "chart", "summarize",
    )
    return sum(1 for verb in action_verbs if verb in lower) >= 2


def _is_modeling_skill(skill: str) -> bool:
    key = (skill or "").lower()
    return any(marker in key for marker in MODELING_SKILL_MARKERS)


def _agent_skill_names(agent: dict) -> list[str]:
    names: list[str] = []
    primary = (agent.get("skill") or "").strip()
    if primary:
        names.append(primary)
    for member in agent.get("member_skills") or []:
        m = (member or "").strip()
        if m and m.lower() not in {n.lower() for n in names}:
            names.append(m)
    return names


def _detect_task_intents(user_message: str) -> set[str]:
    """Infer what kind of work the request needs — independent of static keyword tables."""
    lower = user_message.lower()
    intents: set[str] = set()

    if any(term in lower for term in MACRO_OUTLOOK_TERMS):
        intents.add("macro_outlook")
    if any(term in lower for term in STRATEGY_RECOMMENDATION_TERMS):
        intents.add("strategy_recommendation")
    if any(term in lower for term in DELIVERABLE_REQUEST_TERMS):
        intents.add("stakeholder_delivery")
    if any(term in lower for term in DELIVERABLE_REQUEST_TERMS) and any(
        term in lower for term in EXECUTIVE_WORKFLOW_TERMS
    ):
        intents.add("executive_workflow")
    if any(term in lower for term in DATA_PULL_TERMS):
        intents.add("data_pull")
    if any(
        term in lower
        for term in ("regulation", "regulatory", "compliance", "policy", "legal", "ai act")
    ):
        intents.add("regulatory_research")
    if _requires_quantitative_modeling(user_message):
        intents.add("quant_modeling")
    elif any(term in lower for term in ("forecast", "probability", "outlook", "project")):
        intents.add("qualitative_forecast")
    if _is_research_synthesis_request(user_message):
        intents.add("regulatory_research")

    return intents


def _keyword_hint_score(
    user_message: str,
    agent: dict,
    framework: dict,
) -> int:
    """Fast keyword hints only — accelerates triage but is not authoritative."""
    lower = user_message.lower()
    score = 0

    for skill_name in _agent_skill_names(agent):
        skill_key = skill_name.lower()
        keywords = SKILL_KEYWORDS.get(skill_key, ())
        for kw in keywords:
            if kw in lower:
                score += 2

        tokens = [t for t in re.split(r"[^a-z0-9&]+", skill_key) if len(t) >= 4]
        for tok in tokens:
            if tok in lower:
                score += 1

    for resp in _supports_for_agent(agent, framework):
        resp_lower = resp.lower()
        overlap = sum(1 for w in re.findall(r"[a-z0-9]{5,}", lower) if w in resp_lower)
        score += min(overlap, 2)

    return score


def _semantic_relevance_score(
    user_message: str,
    agent: dict,
) -> int:
    """Score relevance from task intent and specialist role — not exhaustive keywords."""
    skill = (agent.get("skill") or "").lower()
    role = _skill_role(skill)
    score = 0

    for intent in _detect_task_intents(user_message):
        for pattern, weight in INTENT_SKILL_WEIGHTS.get(intent, ()):
            if pattern in skill:
                score += weight

    role_boosts = {
        "interpretation": ("macro_outlook", "strategy_recommendation", "qualitative_forecast"),
        "strategy": ("strategy_recommendation",),
        "review": ("regulatory_research",),
        "delivery": ("stakeholder_delivery", "executive_workflow"),
        "data_gathering": ("data_pull",),
        "analysis": ("quant_modeling", "qualitative_forecast", "executive_workflow"),
    }
    intents = _detect_task_intents(user_message)
    for boosted_role, intent_names in role_boosts.items():
        if role == boosted_role and intents.intersection(intent_names):
            score += 4

    # Penalize delivery/reporting when the user did not ask for a deliverable artifact.
    if role == "delivery" and "stakeholder_delivery" not in intents:
        if intents.intersection({"macro_outlook", "strategy_recommendation", "qualitative_forecast"}):
            score -= 5

    return score


def _rule_relevance_adjustments(
    user_message: str,
    agent: dict,
) -> int:
    """Guardrails on top of semantic + keyword hints."""
    lower = user_message.lower()
    skill = (agent.get("skill") or "").lower()
    score = 0
    research = _is_research_synthesis_request(user_message)
    needs_quant = _requires_quantitative_modeling(user_message)

    if research and not needs_quant:
        if _is_modeling_skill(skill):
            score -= 10
        if any(x in skill for x in ("compliance", "governance", "business acumen", "domain", "investor")):
            score += 4

    if ("invest" in lower or "investment" in lower) and not needs_quant:
        if any(x in skill for x in ("investor", "business acumen", "domain", "compliance")):
            score += 3
        if "valuation" in skill or "dcf" in skill:
            score -= 8

    if needs_quant and _is_modeling_skill(skill):
        score += 4

    return score


def _score_agent_relevance(
    user_message: str,
    agent: dict,
    framework: dict,
) -> int:
    """Combined relevance: semantic intent (primary offline) + keyword hints + guardrails."""
    return (
        _semantic_relevance_score(user_message, agent)
        + _keyword_hint_score(user_message, agent, framework)
        + _rule_relevance_adjustments(user_message, agent)
    )


def _sort_ranked_agents(ranked: list[tuple[dict, float]]) -> list[dict]:
    ranked.sort(
        key=lambda item: (
            -item[1],
            PIPELINE_ORDER.get((item[0].get("skill") or "").lower(), 50),
        ),
    )
    return [agent for agent, _pts in ranked]


def _llm_rank_specialists(
    user_message: str,
    skill_agents: list[dict],
    framework: dict,
    keyword_hints: dict[str, int],
    *,
    progress: ChatProgressReporter | None = None,
) -> list[tuple[dict, float]] | None:
    """Ask the manager LLM which specialists are truly needed (keyword hints are non-binding)."""
    if not cursor_is_configured():
        return None

    lines: list[str] = []
    agents_by_id = {a.get("id"): a for a in skill_agents}
    for agent in skill_agents:
        aid = agent.get("id") or ""
        skill = agent.get("skill") or "Specialist"
        role = _skill_role(skill)
        supports = _supports_for_agent(agent, framework)
        support_preview = "; ".join(supports[:2]) if supports else "general domain work"
        hint = keyword_hints.get(aid, 0)
        hint_note = f" [keyword-hint={hint}]" if hint else ""
        lines.append(
            f"- {aid}: {skill} (role: {role}) — {support_preview}{hint_note}"
        )

    prompt = f"""Rank specialists by how directly they serve the user's request.

User request:
{user_message}

Available specialists (pick only those truly needed — 0 to all, not a default pipeline):
{chr(10).join(lines)}

Keyword hints (optional accelerators only — incomplete and non-binding):
{json.dumps(keyword_hints, indent=2) if keyword_hints else "none"}

Think task-first:
1. What is the core question or decision?
2. Which specialists directly contribute? Skip spreadsheet/valuation/reporting unless explicitly required.

Return JSON only:
{{
  "answer_directly": false,
  "ranked": [
    {{"agent_id": "skill_id", "relevance": 1-10, "reason": "one short sentence"}}
  ]
}}

Rules:
- Macro outlook + positioning/allocation → interpretation/strategy skills, not board reporting unless a report was requested.
- Policy/regulation → compliance/interpretation, not DCF/valuation unless quant was requested.
- Empty ranked list when the manager can answer directly.
- relevance 8-10 = essential, 5-7 = helpful, 1-4 = marginal (omit marginal)."""

    mgr = _manager_agent(framework)
    system = (mgr or {}).get("system_prompt") or "You are a manager agent triaging specialist delegation."
    on_tick = progress.update_planning_thinking if progress else None
    raw = _llm_complete(system, prompt, progress=progress, on_tick=on_tick)
    if not raw:
        return None

    parsed = _parse_ai_json(raw)
    if not parsed:
        return None
    if parsed.get("answer_directly"):
        return []

    ranked: list[tuple[dict, float]] = []
    for entry in parsed.get("ranked") or []:
        aid = entry.get("agent_id")
        agent = agents_by_id.get(aid)
        if not agent:
            continue
        try:
            relevance = float(entry.get("relevance", 0))
        except (TypeError, ValueError):
            relevance = 0.0
        if relevance >= LLM_RELEVANCE_THRESHOLD:
            ranked.append((agent, relevance))
    return ranked


def _rank_relevant_agents(
    user_message: str,
    skill_agents: list[dict],
    framework: dict,
    *,
    min_score: int = 2,
    progress: ChatProgressReporter | None = None,
    use_llm: bool = True,
) -> list[dict]:
    """Rank specialists: LLM triage when available; semantic + keyword hints as fallback."""
    keyword_hints = {
        (agent.get("id") or ""): _keyword_hint_score(user_message, agent, framework)
        for agent in skill_agents
    }

    # LLM triage only during active planning (progress reporter present); tests/offline use semantic.
    if use_llm and progress is not None:
        llm_ranked = _llm_rank_specialists(
            user_message,
            skill_agents,
            framework,
            keyword_hints,
            progress=progress,
        )
        if llm_ranked is not None:
            if not llm_ranked:
                return []
            return _sort_ranked_agents(llm_ranked)

    scored = [
        (agent, _score_agent_relevance(user_message, agent, framework))
        for agent in skill_agents
    ]
    ranked = [(agent, float(pts)) for agent, pts in scored if pts >= min_score]
    if not ranked and scored:
        best = max(scored, key=lambda item: item[1])
        if best[1] > 0:
            ranked = [best]
    return _sort_ranked_agents(ranked)


ROLE_INSTRUCTIONS: dict[str, str] = {
    "data_gathering": (
        "Identify the data, metrics, dimensions, and evidence needed. "
        "Define what to measure and how to pull it — do not interpret or recommend yet."
    ),
    "analysis": (
        "Run analysis on upstream data/findings. Test hypotheses, quantify impact, "
        "and surface patterns relevant to the request."
    ),
    "visualization": (
        "Design charts, dashboards, or visual summaries that communicate the analysis clearly."
    ),
    "interpretation": (
        "Translate technical findings into business meaning, risks, and implications for stakeholders."
    ),
    "strategy": (
        "Recommend strategic options, trade-offs, and prioritized actions based on prior work."
    ),
    "review": (
        "Check outputs for compliance, accuracy, gaps, and alignment with requirements."
    ),
    "delivery": (
        "Produce the final stakeholder-ready deliverable (report section, slides outline, or memo)."
    ),
}

# Skill-specific deliverables shown in delegation plans (distinct per specialist).
SKILL_DELIVERABLES: dict[str, str] = {
    "financial analysis": "Analyze financial performance, KPIs, margins, and trend drivers",
    "financial modeling": "Build or update models, assumptions, sensitivities, and scenarios",
    "valuation": "Estimate valuation range using comps, DCF inputs, and key value drivers",
    "investor reporting": "Prepare investor-facing metrics, narrative, and reporting outputs",
    "catalyst hunter": "Map near-term catalysts, event risks, and timing implications",
    "business acumen": "Frame strategic implications, trade-offs, and executive recommendations",
    "sql query": "Pull, validate, and structure the required dataset and metrics",
    "python automation": "Automate data pulls, transforms, or repeatable analysis steps",
    "statistical modeling": "Run statistical tests, forecasts, or variance decomposition",
    "data visualization": "Design charts or dashboards that communicate findings clearly",
    "presentation & storytelling": "Outline slides and narrative flow for stakeholder delivery",
    "written communication": "Draft the memo, summary, or written deliverable section",
    "growth marketing": "Recommend acquisition, conversion, and channel strategy options",
    "compliance & governance": "Review outputs for policy, accuracy, and compliance gaps",
    "domain expertise": "Apply industry context, benchmarks, and domain-specific judgment",
    "process optimization": "Identify bottlenecks and propose workflow improvements",
    "financial accounting": "Analyze P&L, cash flow, variance, and core accounting outputs",
    "financial modeling & valuation": "Build models, forecasts, and valuation analyses",
    "investor & board reporting": "Prepare investor or board-ready reporting and narrative",
    "data & analytics": "Pull, transform, and analyze the required datasets",
    "statistical analysis": "Run statistical tests, forecasts, or variance decomposition",
}


def _skill_deliverable(skill: str) -> str:
    """One-line scoped deliverable unique to this specialist skill."""
    key = (skill or "").lower().strip()
    for pattern, text in SKILL_DELIVERABLES.items():
        if pattern in key or key in pattern:
            return text
    role = _skill_role(skill)
    role_labels = {
        "data_gathering": f"Gather and validate data for {skill}",
        "analysis": f"Run {skill} analysis on upstream outputs",
        "visualization": f"Visualize results for {skill}",
        "interpretation": f"Interpret findings through a {skill} lens",
        "strategy": f"Recommend actions based on {skill} insights",
        "review": f"Review work through {skill} standards",
        "delivery": f"Produce the {skill} deliverable",
    }
    return role_labels.get(role, f"Complete scoped {skill or 'specialist'} work")


def _subtask_display_summary(st: dict, index: int, total: int) -> str:
    """Human-readable delegation line for plan proposals."""
    skill = st.get("skill") or "Specialist"
    deliverable = _skill_deliverable(skill)
    if index == 0:
        return f"{deliverable} to open the workflow."
    if index == total - 1:
        return f"{deliverable} and synthesize the final answer."
    return f"{deliverable} using upstream outputs (step {index + 1}/{total})."


def _parse_ai_json(text: str) -> dict | None:
    cleaned = text.strip()
    if cleaned.startswith("```"):
        cleaned = re.sub(r"^```(?:json)?\s*", "", cleaned)
        cleaned = re.sub(r"\s*```$", "", cleaned)
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        return None


def _llm_complete(
    system: str,
    user: str,
    *,
    progress: ChatProgressReporter | None = None,
    on_tick=None,
) -> str | None:
    cancel_check = progress.is_cancelled if progress else None
    on_run_start = progress.set_active_run if progress else None
    return cursor_complete(
        system,
        user,
        cancel_check=cancel_check,
        on_tick=on_tick,
        on_run_start=on_run_start,
    )


def _skill_agents(framework: dict) -> list[dict]:
    return [
        a for a in framework.get("agents", [])
        if a.get("type") == "skill" or (
            a.get("id", "").startswith("skill_") and a.get("id") != MANAGER_ID
        )
    ]


def _pipeline_sort(agents: list[dict]) -> list[dict]:
    return sorted(agents, key=lambda a: PIPELINE_ORDER.get((a.get("skill") or "").lower(), 50))


def _manager_agent(framework: dict) -> dict | None:
    for a in framework.get("agents", []):
        if a.get("id") == MANAGER_ID or a.get("type") == "manager":
            return a
    return None


def _skill_role(skill: str) -> str:
    key = (skill or "").lower()
    if key in SKILL_ROLE:
        return SKILL_ROLE[key]
    for pattern, role in SKILL_ROLE.items():
        if pattern in key or key in pattern:
            return role
    return "analysis"


def _supports_for_agent(agent: dict, framework: dict) -> list[str]:
    direct = agent.get("supports_responsibilities") or []
    if direct:
        return direct
    skill = (agent.get("skill") or "").lower()
    for entry in framework.get("skill_breakdown") or []:
        if (entry.get("skill") or "").lower() == skill:
            return entry.get("supports_responsibilities") or []
    return []


def _request_focus(user_message: str) -> str:
    """Short topic label extracted from the user request."""
    text = user_message.strip()
    if len(text) <= 120:
        return text
    return text[:117] + "…"


def _assign_skill_task(
    user_message: str,
    agent: dict,
    index: int,
    total: int,
    framework: dict,
) -> str:
    skill = agent.get("skill") or agent.get("role", "Specialist")
    focus = _request_focus(user_message)
    role = _skill_role(skill)

    if _is_research_synthesis_request(user_message) and not _requires_quantitative_modeling(user_message):
        research_tasks = {
            "compliance & governance": f"Map the regulatory landscape and compliance implications for: «{focus}».",
            "business acumen": f"Assess strategic and investment implications of: «{focus}».",
            "domain expertise": f"Apply sector context and benchmarks relevant to: «{focus}».",
            "investor reporting": f"Frame portfolio/investor implications and narrative for: «{focus}».",
            "written communication": f"Draft a clear synthesis memo addressing: «{focus}».",
            "presentation & storytelling": f"Outline stakeholder-ready narrative for: «{focus}».",
        }
        skill_key = (skill or "").lower()
        for pattern, task in research_tasks.items():
            if pattern in skill_key:
                if index == total - 1 and total > 1:
                    return f"{task} Integrate upstream analysis into the final answer."
                return task

    deliverable = _skill_deliverable(skill)
    if index == 0:
        return f"{deliverable} for: «{focus}»."
    if index == total - 1:
        return f"{deliverable}. Synthesize upstream specialist outputs into the final answer."
    return (
        f"{deliverable}. Continue the workflow (step {index + 1}/{total}) "
        f"using outputs from earlier specialists."
    )


def _pick_default_pipeline(skill_agents: list[dict], max_steps: int = 3) -> list[dict]:
    """Lean chain only when the user explicitly asked for multiple deliverables."""
    ordered = _pipeline_sort(skill_agents)
    picked: list[dict] = []
    for role in ("data_gathering", "analysis", "delivery"):
        for agent in ordered:
            if agent in picked:
                continue
            if _skill_role(agent.get("skill") or "") == role:
                picked.append(agent)
                break
        if len(picked) >= max_steps:
            break
    return picked[:max_steps] if picked else ordered[:max_steps]


def _match_skills_for_request(
    user_message: str,
    skill_agents: list[dict],
    framework: dict | None = None,
    *,
    progress: ChatProgressReporter | None = None,
    use_llm: bool = False,
) -> list[dict]:
    """Return specialists whose skills are plausibly needed for this request."""
    if framework is not None:
        return _rank_relevant_agents(
            user_message,
            skill_agents,
            framework,
            min_score=2,
            progress=progress,
            use_llm=use_llm,
        )
    lower = user_message.lower()
    matched: list[dict] = []
    seen_ids: set[str] = set()
    for agent in skill_agents:
        agent_id = agent.get("id") or ""
        hit = False
        for skill_name in _agent_skill_names(agent):
            skill = skill_name.lower()
            keywords = SKILL_KEYWORDS.get(skill, ())
            tokens = [t for t in re.split(r"[^a-z0-9&]+", skill) if len(t) >= 4]
            if any(kw in lower for kw in keywords) or any(tok in lower for tok in tokens):
                hit = True
                break
        if hit and agent_id not in seen_ids:
            matched.append(agent)
            seen_ids.add(agent_id)
    return _pipeline_sort(matched)


def _fill_specialist_chain(
    seed: list[dict],
    skill_agents: list[dict],
    max_subtasks: int,
) -> list[dict]:
    """Return the seed chain capped at max_subtasks — never pad with unrelated specialists."""
    return _pipeline_sort(seed)[:max_subtasks]


def _complete_deliverable_chain(
    user_message: str,
    ordered: list[dict],
    skill_agents: list[dict],
    framework: dict,
    max_subtasks: int,
    *,
    progress: ChatProgressReporter | None = None,
) -> list[dict]:
    """When the user asked for multiple outputs, add complementary roles from dynamic ranking."""
    if not _explicit_multi_deliverable_request(user_message):
        return ordered[:max_subtasks]
    result = list(ordered)
    target = min(max_subtasks, 3)
    relevance_pool = _rank_relevant_agents(
        user_message, skill_agents, framework, min_score=1, progress=progress, use_llm=False,
    )
    if len(relevance_pool) < 2:
        scored = [
            (agent, float(_score_agent_relevance(user_message, agent, framework)))
            for agent in skill_agents
        ]
        scored = [(agent, pts) for agent, pts in scored if pts >= 1]
        if scored:
            relevance_pool = _sort_ranked_agents(scored)
    for role in ("data_gathering", "analysis", "delivery"):
        if len(result) >= target:
            break
        for agent in relevance_pool:
            if agent in result:
                continue
            if _skill_role(agent.get("skill") or "") != role:
                continue
            if _score_agent_relevance(user_message, agent, framework) < 1:
                continue
            result.append(agent)
            break
    return result[:max_subtasks]


def _manager_decompose_rule(
    user_message: str,
    agent_name: str,
    framework: dict,
    skill_agents: list[dict],
    preferred_agents: list[dict] | None = None,
    max_subtasks: int = DEFAULT_MAX_SUBTASKS,
    progress: ChatProgressReporter | None = None,
) -> dict:
    if not skill_agents:
        return {
            "reasoning": f"{agent_name} handles this directly (no skill team configured).",
            "subtasks": [],
        }

    ranked = _rank_relevant_agents(
        user_message, skill_agents, framework, min_score=2, progress=progress,
    )
    if preferred_agents:
        preferred_ids = {a.get("id") for a in preferred_agents}
        ordered = [a for a in ranked if a.get("id") in preferred_ids]
        if not ordered:
            ordered = ranked[:max_subtasks] or _pipeline_sort(preferred_agents)[:max_subtasks]
    elif ranked:
        ordered = ranked[:max_subtasks]
    else:
        return {
            "reasoning": (
                f"{agent_name} can answer this directly — dynamic triage found no specialists "
                f"whose skills directly serve this request."
            ),
            "subtasks": [],
        }

    ordered = _complete_deliverable_chain(
        user_message, ordered, skill_agents, framework, max_subtasks, progress=progress,
    )

    subtasks = []
    for i, agent in enumerate(ordered):
        skill = agent.get("skill") or agent.get("role", "Specialist")
        subtasks.append({
            "skill": skill,
            "agent_id": agent.get("id"),
            "task": _assign_skill_task(user_message, agent, i, len(ordered), framework),
        })

    skill_chain = " → ".join(st["skill"] for st in subtasks)
    skipped = len(skill_agents) - len(ordered)
    reasoning = (
        f"Manager plan for {agent_name}: delegate to {len(subtasks)} specialist"
        f"{'s' if len(subtasks) != 1 else ''} ({skill_chain}). "
        f"Skipped {skipped} specialist{'s' if skipped != 1 else ''} not required for this question."
    )
    return {"reasoning": reasoning, "subtasks": subtasks}


def _resolve_subtask_agent(
    subtask: dict,
    skill_agents: list[dict],
) -> dict | None:
    agent_id = subtask.get("agent_id")
    if agent_id:
        for agent in skill_agents:
            if agent.get("id") == agent_id:
                return agent
    skill = (subtask.get("skill") or "").lower()
    for agent in skill_agents:
        agent_skill = (agent.get("skill") or "").lower()
        if agent_skill == skill or skill in agent_skill or agent_skill in skill:
            return agent
    return None


def _sanitize_manager_plan(
    plan: dict,
    user_message: str,
    framework: dict,
    skill_agents: list[dict],
    max_subtasks: int,
) -> dict:
    """Drop irrelevant specialists the manager may have over-assigned."""
    sanitized: list[dict] = []
    needs_quant = _requires_quantitative_modeling(user_message)
    research = _is_research_synthesis_request(user_message)

    for st in plan.get("subtasks") or []:
        agent = _resolve_subtask_agent(st, skill_agents)
        if not agent:
            continue
        skill = agent.get("skill") or st.get("skill") or "Specialist"
        if research and not needs_quant and _is_modeling_skill(skill):
            continue
        score = _score_agent_relevance(user_message, agent, framework)
        if score < 1 and len(sanitized) >= 1:
            continue
        sanitized.append({
            "skill": skill,
            "agent_id": agent.get("id"),
            "task": (st.get("task") or "").strip() or _assign_skill_task(
                user_message, agent, len(sanitized), max(len(sanitized) + 1, 1), framework,
            ),
        })
        if len(sanitized) >= max_subtasks:
            break

    if not sanitized:
        return {
            "reasoning": plan.get("reasoning") or f"Answering directly — no specialist workflow needed.",
            "subtasks": [],
        }

    skill_chain = " → ".join(st["skill"] for st in sanitized)
    reasoning = plan.get("reasoning") or (
        f"Delegated to {len(sanitized)} specialist{'s' if len(sanitized) != 1 else ''} ({skill_chain})."
    )
    return {"reasoning": reasoning, "subtasks": sanitized}


def _manager_decompose_ai(
    user_message: str,
    agent_name: str,
    framework: dict,
    skill_agents: list[dict],
    progress: ChatProgressReporter | None = None,
    preferred_agents: list[dict] | None = None,
    max_subtasks: int = DEFAULT_MAX_SUBTASKS,
    agent_context: dict | None = None,
    memory_context: ConversationMemoryContext | None = None,
) -> dict | None:
    agents_pool = skill_agents
    agents_desc_lines = []
    for a in agents_pool[:10]:
        supports = _supports_for_agent(a, framework) or []
        support_preview = "; ".join(supports[:2]) if supports else "general domain work"
        agents_desc_lines.append(
            f"- {a.get('id')}: {a.get('skill')} — typically handles: {support_preview}"
        )
    agents_desc = "\n".join(agents_desc_lines)

    ranked = _rank_relevant_agents(
        user_message, skill_agents, framework, min_score=2, progress=progress,
    )
    hint = ""
    if ranked:
        names = ", ".join(a.get("skill") or "?" for a in ranked[:4])
        hint = (
            f"\nRelevance scan (task-first, not roster-wide): {names} may be useful. "
            f"Still skip any skill that does not directly serve the user's question.\n"
        )
    if _is_research_synthesis_request(user_message) and not _requires_quantitative_modeling(user_message):
        hint += (
            "\nThis looks like a research/synthesis question — do NOT default to valuation, DCF, "
            "financial modeling, cash flow, or forecasting unless the user explicitly asked for quant work.\n"
        )

    memory_section = ""
    if memory_context and memory_context.memory_prompt_block:
        memory_section = f"\n{memory_context.memory_prompt_block}\n"

    prompt = f"""You are the Manager Agent for {agent_name}.

User request: {user_message}
{memory_section}
Available specialists (do NOT use all of them — pick only what the question needs):
{agents_desc}
{hint}
Think task-first:

1. State the user's core question in one sentence (what decision or insight they need).
2. List the concrete deliverables required to answer it (research, analysis, memo, model, etc.).
3. Map ONLY those deliverables to specialists. If none are needed, answer directly yourself.

Rules:
- Return an EMPTY subtasks list when you can answer directly (definitions, greetings, simple facts).
- Do NOT run the full finance/data pipeline by default — skip specialists whose skills are irrelevant.
- Do NOT assign valuation/DCF/modeling/cash-flow/forecasting for policy, macro, or general investment-impact questions unless quant output was explicitly requested.
- Use at most {max_subtasks} sequential subtasks. Each subtask must:
  - Map to exactly one specialist
  - Be a DISTINCT scoped assignment (not a copy of the user message)
  - Build on earlier steps without repeating work

Return JSON only:
{{
  "reasoning": "2-3 sentences: core question + why these specialists (or direct reply)",
  "subtasks": [
    {{"skill": "exact skill name", "agent_id": "skill_id", "task": "specific scoped assignment for this specialist only"}}
  ]
}}"""
    mgr = _manager_agent(framework)
    system = (mgr or {}).get("system_prompt") or f"You are the Manager Agent for {agent_name}."
    system = _compose_system(system, agent_context, memory_context)
    on_tick = progress.update_planning_thinking if progress else None
    raw = _llm_complete(system, prompt, progress=progress, on_tick=on_tick)
    if not raw:
        return None
    parsed = _parse_ai_json(raw)
    if not parsed or "subtasks" not in parsed:
        return None
    subtasks = parsed.get("subtasks") or []
    if not subtasks:
        return {
            "reasoning": parsed.get("reasoning") or f"{agent_name} will answer directly.",
            "subtasks": [],
        }
    plan = {
        "reasoning": parsed.get("reasoning") or "Manager decomposed the request into specialist tasks.",
        "subtasks": subtasks[:max_subtasks],
    }
    return _sanitize_manager_plan(plan, user_message, framework, skill_agents, max_subtasks)


def _apply_plan_feedback(user_message: str, user_feedback: str | None) -> str:
    if not user_feedback or not user_feedback.strip():
        return user_message
    return (
        f"{user_message}\n\n"
        f"User revision notes (adjust the delegation plan accordingly):\n{user_feedback.strip()}"
    )


def plan_requires_confirmation(subtasks: list[dict], skill_agents: list[dict]) -> bool:
    """True when the plan delegates to more than half of available specialists."""
    if not subtasks or len(skill_agents) < 2:
        return False
    return len(subtasks) > len(skill_agents) / 2


def format_plan_proposal_message(agent_name: str, plan: dict, skill_agents: list[dict]) -> str:
    subtasks = plan.get("subtasks") or []
    lines = [
        f"I've drafted a workflow that uses **{len(subtasks)} of {len(skill_agents)}** specialists "
        f"on my team — more than half — so I'd like your approval before starting.",
        "",
        f"**Manager plan:** {plan.get('reasoning') or 'Delegated specialist workflow.'}",
        "",
        "**Delegation steps:**",
    ]
    for i, st in enumerate(subtasks, 1):
        skill = st.get("skill") or "Specialist"
        summary = _subtask_display_summary(st, i - 1, len(subtasks))
        lines.append(f"{i}. **{skill}** — {summary}")
    lines.extend([
        "",
        "Review the workflow below. **Confirm** to start execution, or add comments to revise the plan.",
    ])
    return "\n".join(lines)


def _manager_decompose(
    user_message: str,
    agent_name: str,
    framework: dict,
    skill_agents: list[dict],
    progress: ChatProgressReporter | None = None,
    preferred_agents: list[dict] | None = None,
    user_feedback: str | None = None,
    max_subtasks: int = DEFAULT_MAX_SUBTASKS,
    agent_context: dict | None = None,
    memory_context: ConversationMemoryContext | None = None,
) -> dict:
    prompt_message = _apply_plan_feedback(user_message, user_feedback)
    rule_plan = _manager_decompose_rule(
        prompt_message,
        agent_name,
        framework,
        skill_agents,
        preferred_agents,
        max_subtasks,
        progress=progress,
    )
    if rule_plan.get("subtasks"):
        return rule_plan

    ai_plan = _manager_decompose_ai(
        prompt_message,
        agent_name,
        framework,
        skill_agents,
        progress,
        preferred_agents,
        max_subtasks,
        agent_context,
        memory_context,
    )
    if ai_plan is None:
        return rule_plan

    return ai_plan


def _rule_subtask_output(
    subtask: dict,
    agent: dict,
    agent_name: str,
    prior_outputs: str,
) -> str:
    skill = subtask.get("skill") or agent.get("skill", "Specialist")
    task = subtask.get("task") or ""
    role = _skill_role(skill)
    prior_note = ""
    if prior_outputs.strip():
        prior_note = f"\n\n**Building on upstream work:**\n{prior_outputs.strip()[-1200:]}"

    skill_lower = skill.lower()
    if role == "data_gathering" or "sql" in skill_lower:
        body = (
            f"- Metrics and dimensions to pull\n"
            f"- Suggested tables/filters and aggregation logic\n"
            f"- Data quality checks before analysis"
        )
    elif role == "visualization":
        body = (
            f"- Recommended chart types and axes\n"
            f"- KPIs to highlight\n"
            f"- Annotations/callouts for stakeholders"
        )
    elif role in ("delivery",) or "communication" in skill_lower or "presentation" in skill_lower:
        body = (
            f"- Executive summary bullets\n"
            f"- Recommended report/deck structure\n"
            f"- Key message and call-to-action"
        )
    elif role == "interpretation" or "business" in skill_lower:
        body = (
            f"- Business interpretation of findings\n"
            f"- Risks, assumptions, and confidence level\n"
            f"- Stakeholder implications"
        )
    else:
        body = (
            f"- Analytical approach for this step\n"
            f"- Key findings from {skill.lower()} perspective\n"
            f"- Outputs to pass downstream"
        )

    return (
        f"**{skill} specialist output**\n\n"
        f"**Assignment:** {task}\n\n"
        f"{body}"
        f"{prior_note}"
    )


def _run_subtask(
    subtask: dict,
    agent: dict,
    agent_name: str,
    prior_outputs: str,
    file_context: str,
    progress: ChatProgressReporter | None = None,
    agent_context: dict | None = None,
    memory_context: ConversationMemoryContext | None = None,
) -> str:
    skill = subtask.get("skill") or agent.get("skill", "Specialist")
    assignment = subtask.get("task") or f"Apply {skill} to the manager's current workflow."
    system = agent.get("system_prompt") or f"You are {agent_name}'s {skill} specialist."
    skill_md = agent.get("skill_md") or ""

    user = f"""MANAGER ASSIGNMENT (your only job — do not answer the full user request yourself):
{assignment}

Rules:
- Complete ONLY the work described in the manager assignment above.
- Do NOT repeat work assigned to other specialists.
- If upstream specialist outputs are provided, use them as inputs — do not redo their analysis.
{f"Upstream specialist outputs:{chr(10)}{prior_outputs}" if prior_outputs else ""}
{f"Reference file context (use if relevant to your assignment):{chr(10)}{file_context[:4000]}" if file_context else ""}

Provide focused, actionable output for the {skill} step only."""

    if skill_md:
        system = f"{system}\n\n--- skill.md ---\n{skill_md[:3000]}"
    system = _compose_system(system, agent_context, memory_context)

    if progress:
        progress.raise_if_cancelled()
    on_tick = progress.update_delegation_thinking if progress else None
    ai = _llm_complete(system, user, progress=progress, on_tick=on_tick)
    if progress and progress.is_cancelled():
        raise ChatGenerationCancelled()
    if ai:
        return ai.strip()

    return _rule_subtask_output(subtask, agent, agent_name, prior_outputs)


def _consolidate(
    agent_name: str,
    user_message: str,
    subtask_results: list[dict],
    manager: dict | None,
    manager_reasoning: str,
    progress: ChatProgressReporter | None = None,
    agent_context: dict | None = None,
    memory_context: ConversationMemoryContext | None = None,
) -> str:
    combined = "\n\n".join(
        f"**{r['skill']}:**\n{r['output']}" for r in subtask_results
    )
    mgr_prompt = (manager or {}).get("system_prompt") or f"You are the manager for {agent_name}."
    user = f"""The user asked: {user_message}

Your decomposition plan: {manager_reasoning}

Specialist outputs (already completed — synthesize, do not re-list raw subtasks):
{combined}

Write ONE consolidated reply as {agent_name} — first person, professional, no mention of sub-agents or internal steps.
Include key findings and next steps."""
    mgr_prompt = _compose_system(mgr_prompt, agent_context, memory_context)
    if progress:
        progress.raise_if_cancelled()
    on_tick = progress.update_delegation_thinking if progress else None
    ai = _llm_complete(mgr_prompt, user, progress=progress, on_tick=on_tick)
    if progress and progress.is_cancelled():
        raise ChatGenerationCancelled()
    if ai:
        return ai.strip()
    return (
        f"Here's my take on your request:\n\n{combined}\n\n"
        f"Let me know if you'd like me to go deeper on any area."
    )


def _trivial_message(text: str) -> bool:
    normalized = re.sub(r"[^\w\s]", "", text.lower()).strip()
    words = normalized.split()
    if not words:
        return True
    if len(words) <= 3 and normalized in TRIVIAL_MESSAGES:
        return True
    return len(words) <= 2 and all(w in TRIVIAL_MESSAGES for w in words)


def _is_direct_answer(user_message: str) -> bool:
    """Fast path: manager can answer without any specialist delegation."""
    text = normalize_user_query(user_message)
    if _trivial_message(text):
        return True
    lower = text.lower()
    if any(re.search(pat, lower) for pat in DIRECT_ANSWER_PATTERNS):
        return True
    word_count = len(text.split())
    if word_count <= 12 and re.match(r"^(what|who|when|where|how) (is|are|was|were)\b", lower):
        if not any(kw in lower for kw in COMPLEX_KEYWORDS):
            return True
    if try_instant_reply(user_message, "Agent"):
        return True
    return False


def _is_multi_step_request(user_message: str) -> bool:
    """True when the request likely needs multiple specialists in sequence."""
    lower = user_message.lower()
    if any(kw in lower for kw in COMPLEX_KEYWORDS):
        if any(w in lower for w in (" and ", " then ", "report", "prepare", "recommend", "summary")):
            return True
        if len(user_message.split()) >= 12:
            return True
    return len(user_message.split()) >= 20


def _should_use_team_mode(
    user_message: str,
    skill_agents: list[dict],
    framework: dict | None = None,
) -> bool:
    if len(skill_agents) < 2:
        return False
    if _is_direct_answer(user_message):
        return False
    if _explicit_multi_deliverable_request(user_message):
        return True
    matched = (
        _match_skills_for_request(user_message, skill_agents, framework)
        if framework is not None
        else _match_skills_for_request(user_message, skill_agents)
    )
    if len(matched) >= 2:
        return True
    if _is_multi_step_request(user_message):
        return True
    return False


def _direct_reply_meta(agent_name: str, cloud_agent_id: str | None) -> dict:
    return {
        "mode": "direct",
        "agent_name": agent_name,
        "cursor_cloud_agent_id": cloud_agent_id,
    }


def _simple_reply_meta(agent_name: str, cloud_agent_id: str | None) -> dict:
    return {
        "mode": "simple",
        "agent_name": agent_name,
        "cursor_cloud_agent_id": cloud_agent_id,
    }


def _simple_reply(
    agent_name: str,
    framework: dict,
    user_message: str,
    history: list[dict],
    file_context: str,
    cloud_agent_id: str | None = None,
    progress: ChatProgressReporter | None = None,
    system_override: str | None = None,
    agent_context: dict | None = None,
    memory_context: ConversationMemoryContext | None = None,
) -> tuple[str, str | None]:
    instant = try_instant_reply(user_message, agent_name)
    if instant:
        return instant, cloud_agent_id

    if system_override:
        system = system_override
    else:
        manager = _manager_agent(framework)
        system = (manager or {}).get("system_prompt") or f"You are {agent_name}, a professional AI agent."
        system = f"You are {agent_name}. Reply in first person as {agent_name}. Be helpful and concise.\n{system}"

    system = _compose_system(system, agent_context, memory_context)
    prior = _effective_history(history, memory_context)
    llm_user_message = _wrap_user_message_for_follow_up(
        user_message,
        cloud_agent_id,
        agent_context,
        memory_context,
        recent_history=prior,
    )
    cancel_check = progress.is_cancelled if progress else None
    on_tick = progress.update_simple_thinking if progress else None
    on_run_start = progress.set_active_run if progress else None

    def _call(cloud_id: str | None, with_history: list[dict] | None) -> tuple[str | None, str | None]:
        return cursor_chat(
            system,
            llm_user_message if cloud_id else user_message,
            history=with_history,
            cloud_agent_id=cloud_id,
            agent_name=f"Minion chat — {agent_name}",
            file_context=file_context,
            cancel_check=cancel_check,
            on_tick=on_tick,
            on_run_start=on_run_start,
        )

    text, new_agent_id = _call(
        cloud_agent_id,
        prior if not cloud_agent_id else None,
    )
    if progress and progress.is_cancelled():
        raise ChatGenerationCancelled()

    if not text and cloud_agent_id:
        text, new_agent_id = _call(None, prior)

    if progress and progress.is_cancelled():
        raise ChatGenerationCancelled()
    if text:
        return text.strip(), new_agent_id or cloud_agent_id

    if progress and progress.is_cancelled():
        raise ChatGenerationCancelled()

    return (
        f"Hi, I'm {agent_name}. I received your message: \"{user_message[:200]}\". "
        f"How can I help you further?",
        cloud_agent_id,
    )


def _generate_docx_artifact(
    user_id: int,
    agent_session_id: int,
    agent_name: str,
    title: str,
    body: str,
) -> dict | None:
    try:
        from docx import Document
    except ImportError:
        return None

    slug = agent_file_slug(agent_name)
    safe_title = re.sub(r"[^a-z0-9]+", "_", title.lower()).strip("_")[:40] or "deliverable"
    filename = f"{slug}_{safe_title}.docx"

    out_dir = AGENT_OUTPUT_FOLDER / str(user_id) / f"chat-{agent_session_id}"
    out_dir.mkdir(parents=True, exist_ok=True)
    path = out_dir / filename

    doc = Document()
    doc.add_heading(title, 0)
    for para in body.split("\n\n"):
        if para.strip():
            doc.add_paragraph(para.strip())
    doc.save(str(path))

    return {
        "filename": filename,
        "path": str(path),
        "artifact_type": "chat_deliverable_docx",
        "size_bytes": path.stat().st_size,
        "title": title,
    }


def _maybe_deliverable(user_message: str, consolidated: str, agent_name: str) -> dict | None:
    lower = user_message.lower()
    if not any(w in lower for w in ("report", "deck", "summary", "document", "write up", "write-up")):
        return None
    title = f"Summary for {agent_name}"
    if "migration" in lower:
        title = "Migration Analysis Summary"
    elif "funnel" in lower:
        title = "Funnel Analysis Summary"
    return {"title": title, "body": consolidated}


def _specialist_reply(
    agent_name: str,
    agent: dict,
    user_message: str,
    history: list[dict],
    file_context: str,
    cloud_agent_id: str | None = None,
    progress: ChatProgressReporter | None = None,
    agent_context: dict | None = None,
    memory_context: ConversationMemoryContext | None = None,
) -> tuple[str, str | None]:
    """Single specialist answers directly — one LLM call, no manager consolidation."""
    skill = agent.get("skill") or agent.get("role", "Specialist")
    system = agent.get("system_prompt") or f"You are {agent_name}'s {skill} specialist."
    skill_md = agent.get("skill_md") or ""
    if skill_md:
        system = f"{system}\n\n--- skill.md ---\n{skill_md[:3000]}"
    system = (
        f"You are {agent_name}. Reply in first person as {agent_name}, "
        f"applying your {skill} expertise. Be helpful and concise.\n{system}"
    )
    framework = {"agents": [agent]}
    return _simple_reply(
        agent_name,
        framework,
        user_message,
        history,
        file_context,
        cloud_agent_id,
        progress,
        system_override=system,
        agent_context=agent_context,
        memory_context=memory_context,
    )


def _execute_team_plan(
    plan: dict,
    *,
    context: dict,
    framework: dict,
    user_message: str,
    skill_agents: list[dict],
    manager: dict | None,
    agents_by_id: dict,
    file_context: str,
    user_id: int | None,
    agent_session_id: int | None,
    cloud_agent_id: str | None,
    progress: ChatProgressReporter | None,
    memory_context: ConversationMemoryContext | None = None,
) -> dict:
    agent_name = context.get("full_name") or "Agent"
    subtasks = plan.get("subtasks") or []
    manager_reasoning = plan.get("reasoning") or ""

    if progress:
        progress.set_manager_plan(manager_reasoning, subtasks, skill_agents)

    progress_subtasks = []
    subtask_results = []
    prior = ""

    for i, st in enumerate(subtasks):
        if progress:
            progress.raise_if_cancelled()
            progress.step_active(i, st)
        agent_id = st.get("agent_id")
        agent = agents_by_id.get(agent_id) or next(
            (a for a in skill_agents if (a.get("skill") or "").lower() == (st.get("skill") or "").lower()),
            skill_agents[min(i, len(skill_agents) - 1)] if skill_agents else {},
        )
        skill = st.get("skill") or agent.get("skill", "Specialist")
        task_label = st.get("task") or skill
        progress_subtasks.append({
            "index": i + 1,
            "skill": skill,
            "agent_id": agent.get("id"),
            "status": "in_progress",
            "label": task_label[:120],
        })

        output = _run_subtask(
            st,
            agent,
            agent_name,
            prior,
            file_context,
            progress,
            agent_context=context,
            memory_context=memory_context,
        )
        progress_subtasks[-1]["status"] = "done"
        if progress:
            progress.step_done(i)
        subtask_results.append({"skill": skill, "output": output, "task": task_label})
        prior += f"\n\n[{skill}]\n{output}"

    if len(subtask_results) == 1:
        content = subtask_results[0]["output"]
    else:
        if progress:
            progress.begin_synthesizing()
        content = _consolidate(
            agent_name,
            user_message,
            subtask_results,
            manager,
            manager_reasoning,
            progress,
            agent_context=context,
            memory_context=memory_context,
        )

    progress_card = {
        "agent_name": agent_name,
        "status": "completed",
        "manager_plan": manager_reasoning,
        "subtasks": progress_subtasks,
        "total": len(progress_subtasks),
        "completed": len(progress_subtasks),
        "summary": f"{agent_name}'s team completed {len(progress_subtasks)} specialist steps",
    }

    meta: dict = {
        "mode": "team_task",
        "agent_name": agent_name,
        "progress_card": progress_card,
        "artifacts": [],
        "cursor_cloud_agent_id": cloud_agent_id,
    }

    deliverable = _maybe_deliverable(user_message, content, agent_name)
    if deliverable and user_id and agent_session_id:
        artifact = _generate_docx_artifact(
            user_id,
            agent_session_id,
            agent_name,
            deliverable["title"],
            deliverable["body"],
        )
        if artifact:
            meta["artifacts"].append(artifact)

    return {"content": content, "meta": meta}


def run_agent_turn(
    context: dict,
    framework: dict,
    user_message: str,
    history: list[dict] | None = None,
    file_context: str = "",
    user_id: int | None = None,
    agent_session_id: int | None = None,
    cloud_agent_id: str | None = None,
    progress: ChatProgressReporter | None = None,
    approved_plan: dict | None = None,
    user_feedback: str | None = None,
    memory_context: ConversationMemoryContext | None = None,
) -> dict:
    """
    Process one user message. Returns assistant reply + metadata (progress_card, artifacts).
    """
    history = history or []
    agent_name = context.get("full_name") or "Agent"
    skill_agents = _skill_agents(framework)
    manager = _manager_agent(framework)
    agents_by_id = {a.get("id"): a for a in framework.get("agents", [])}

    if progress:
        progress.raise_if_cancelled()

    if approved_plan and approved_plan.get("subtasks"):
        return _execute_team_plan(
            approved_plan,
            context=context,
            framework=framework,
            user_message=user_message,
            skill_agents=skill_agents,
            manager=manager,
            agents_by_id=agents_by_id,
            file_context=file_context,
            user_id=user_id,
            agent_session_id=agent_session_id,
            cloud_agent_id=cloud_agent_id,
            progress=progress,
            memory_context=memory_context,
        )

    if _is_direct_answer(user_message) or len(skill_agents) < 2:
        if progress:
            progress.begin_simple(
                agent_name,
                user_message,
                routing="direct" if _is_direct_answer(user_message) else "simple",
            )
        content, updated_cloud_agent_id = _simple_reply(
            agent_name,
            framework,
            user_message,
            history,
            file_context,
            cloud_agent_id,
            progress,
            agent_context=context,
            memory_context=memory_context,
        )
        meta = _direct_reply_meta(agent_name, updated_cloud_agent_id)
        if not _is_direct_answer(user_message):
            meta["mode"] = "simple"
        return {"content": content, "meta": meta}

    matched_skills = _match_skills_for_request(
        user_message, skill_agents, framework, use_llm=False,
    )
    if len(matched_skills) == 1 and not _is_multi_step_request(user_message):
        if progress:
            progress.begin_simple(
                agent_name,
                user_message,
                routing="specialist",
                skill=matched_skills[0].get("skill"),
            )
        content, updated_cloud_agent_id = _specialist_reply(
            agent_name,
            matched_skills[0],
            user_message,
            history,
            file_context,
            cloud_agent_id,
            progress,
            agent_context=context,
            memory_context=memory_context,
        )
        return {
            "content": content,
            "meta": {
                "mode": "specialist",
                "agent_name": agent_name,
                "skill": matched_skills[0].get("skill"),
                "cursor_cloud_agent_id": updated_cloud_agent_id,
            },
        }

    if not _should_use_team_mode(user_message, skill_agents, framework):
        if progress:
            progress.begin_simple(agent_name, user_message)
        content, updated_cloud_agent_id = _simple_reply(
            agent_name,
            framework,
            user_message,
            history,
            file_context,
            cloud_agent_id,
            progress,
            agent_context=context,
            memory_context=memory_context,
        )
        return {"content": content, "meta": _simple_reply_meta(agent_name, updated_cloud_agent_id)}

    if _is_multi_step_request(user_message):
        preferred = matched_skills if len(matched_skills) >= 2 else None
    else:
        preferred = matched_skills or None
    if progress:
        progress.begin_planning(
            agent_name,
            user_message,
            skill_agents,
            matched_skills=matched_skills or preferred,
        )
    expanded_planning = _use_expanded_planning_cap(
        user_message, skill_agents, matched_skills, user_feedback
    )
    max_subtasks = _max_subtasks(skill_agents, expanded=expanded_planning)
    plan = _manager_decompose(
        user_message,
        agent_name,
        framework,
        skill_agents,
        progress,
        preferred,
        user_feedback,
        max_subtasks,
        context,
        memory_context,
    )
    subtasks = plan.get("subtasks") or []

    if not subtasks:
        if progress:
            progress.begin_simple(agent_name, user_message, routing="direct")
        content, updated_cloud_agent_id = _simple_reply(
            agent_name,
            framework,
            user_message,
            history,
            file_context,
            cloud_agent_id,
            progress,
            agent_context=context,
            memory_context=memory_context,
        )
        return {"content": content, "meta": _direct_reply_meta(agent_name, updated_cloud_agent_id)}

    if plan_requires_confirmation(subtasks, skill_agents):
        if progress:
            progress.set_manager_plan(plan.get("reasoning") or "", subtasks, skill_agents)
        proposal_steps = [
            {
                "index": i + 1,
                "skill": st.get("skill") or "Specialist",
                "status": "queued",
                "label": _subtask_display_summary(st, i, len(subtasks)),
            }
            for i, st in enumerate(subtasks)
        ]
        return {
            "needs_confirmation": True,
            "content": format_plan_proposal_message(agent_name, plan, skill_agents),
            "meta": {
                "type": "plan_proposal",
                "mode": "plan_proposal",
                "agent_name": agent_name,
                "plan": plan,
                "skill_count": len(skill_agents),
                "delegated_count": len(subtasks),
                "progress_card": {
                    "agent_name": agent_name,
                    "status": "awaiting_confirmation",
                    "manager_plan": plan.get("reasoning") or "",
                    "subtasks": proposal_steps,
                    "total": len(proposal_steps),
                    "completed": 0,
                    "summary": f"Awaiting approval — {len(subtasks)} specialist steps proposed",
                },
            },
            "plan": plan,
        }

    return _execute_team_plan(
        plan,
        context=context,
        framework=framework,
        user_message=user_message,
        skill_agents=skill_agents,
        manager=manager,
        agents_by_id=agents_by_id,
        file_context=file_context,
        user_id=user_id,
        agent_session_id=agent_session_id,
        cloud_agent_id=cloud_agent_id,
        progress=progress,
        memory_context=memory_context,
    )


def welcome_message(context: dict, framework: dict) -> str:
    agent_name = context.get("full_name") or "your agent"
    skills = [
        a.get("skill") or a.get("role")
        for a in framework.get("agents", [])
        if a.get("type") == "skill" or str(a.get("id", "")).startswith("skill_")
    ]
    skill_line = ", ".join(skills[:5]) if skills else "several specialists"
    job = context.get("current_job") or context.get("field") or "your field"
    return (
        f"Hi — I'm **{agent_name}**, your {job} agent. Behind the scenes I work with a team "
        f"({skill_line}) to handle complex tasks.\n\n"
        f"Ask me anything, or try:\n"
        f"- \"Analyze this data and summarize key findings\"\n"
        f"- \"Prepare a report on [topic]\"\n\n"
        f"Mention **@{agent_name}** in your message, or just type your question."
    )
